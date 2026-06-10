import os
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
import tenacity

from src.legal_ai.rag import retrieve, documents, filenames, embeddings
from src.legal_ai.config import config
from src.legal_ai.logger import setup_logging, get_logger

# ======================================
# CONFIGURE LOGGING
# ======================================

setup_logging(config.LOGS_DIR, config.DEBUG)
logger = get_logger(__name__)

# ======================================
# CONFIGURE GEMINI WITH RETRY LOGIC
# ======================================

try:
    genai.configure(api_key=config.GEMINI_API_KEY)
    model = genai.GenerativeModel(config.GEMINI_MODEL)
    logger.info(f"Successfully configured Gemini model: {config.GEMINI_MODEL}")
except Exception as e:
    logger.error(f"Failed to configure Gemini: {e}")
    raise


@tenacity.retry(
    wait=tenacity.wait_exponential(multiplier=1, min=4, max=10),
    stop=tenacity.stop_after_attempt(3),
    reraise=True
)
def call_gemini(prompt: str, temperature: float = 0.2, max_tokens: int = 2048) -> str:
    """
    Call Gemini API with retry logic.
    Retries up to 3 times with exponential backoff on failure.
    """
    try:
        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": temperature,
                "max_output_tokens": max_tokens
            }
        )
        return response.text
    except Exception as e:
        logger.error(f"Gemini API call failed: {e}")
        raise


def validate_input(user_input: str, max_length: int = 5000) -> bool:
    """Validate user input."""
    if not user_input or not user_input.strip():
        logger.warning("Empty user input provided")
        return False
    
    if len(user_input) > max_length:
        logger.warning(f"User input exceeds max length of {max_length} characters")
        return False
    
    return True


def save_contract(
    contract_text: str,
    output_folder: str = "outputs",
    contract_type: str = "legal_agreement"
) -> str:
    """
    Save contract to DOCX file with unique timestamp.
    
    Returns:
        Path to saved file
    """
    try:
        output_path = config.OUTPUTS_DIR / f"{contract_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc = Document()
        doc.add_heading("Generated Legal Agreement", level=1)
        doc.add_heading(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", level=3)
        
        # Add paragraphs to preserve formatting
        for paragraph_text in contract_text.split('\n\n'):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text)
        
        doc.save(output_path)
        logger.info(f"Contract saved to: {output_path}")
        
        return str(output_path)
    
    except Exception as e:
        logger.error(f"Failed to save contract: {e}")
        raise


# ======================================
# MAIN APPLICATION
# ======================================

def main():
    print("\n" + "="*40)
    print(" AI LEGAL CONTRACT GENERATOR")
    print("="*40 + "\n")
    
    # Check if documents are loaded
    if not documents:
        logger.warning("No contract templates loaded. Proceeding with generation only.")
    
    # Get user input
    user_request = input("Describe the contract you want:\n\n> ").strip()
    
    if not validate_input(user_request):
        logger.error("Invalid user input")
        print("ERROR: Invalid input provided")
        return
    
    logger.info(f"User request received: {user_request[:100]}...")
    
    try:
        # ======================================
        # STEP 1 — STRUCTURED LEGAL SPECIFICATION
        # ======================================
        
        print("\n[1/5] Creating structured legal specification...")
        logger.info("Step 1: Creating structured legal specification")
        
        spec_prompt = f"""You are a senior legal solutions architect.

A user has provided a rough request for a legal agreement.
Transform it into a professional structured legal drafting specification.

USER REQUEST:
{user_request}

Create output using EXACTLY this structure:

CONTRACT TYPE:
...

JURISDICTION:
...

PARTIES:
...

BUSINESS MODEL:
...

COMMERCIAL TERMS:
...

RISK POSITION:
...

REQUIRED CLAUSES:
...

SPECIAL REQUIREMENTS:
...

DRAFTING INSTRUCTIONS:
...

If information is missing: make reasonable assumptions and clearly state them."""
        
        structured_spec = call_gemini(spec_prompt, temperature=0.2, max_tokens=2048)
        
        print("\n" + "="*40)
        print(" STRUCTURED LEGAL SPECIFICATION")
        print("="*40)
        print(structured_spec)
        
        # ======================================
        # STEP 2 — RETRIEVE RELEVANT CLAUSES
        # ======================================
        
        print("\n[2/5] Retrieving relevant contract templates...")
        logger.info("Step 2: Retrieving relevant clauses")
        
        results = retrieve(
            structured_spec,
            documents,
            filenames,
            embeddings,
            top_k=config.RETRIEVAL_TOP_K
        )
        
        context = ""
        if results:
            for score, doc, filename in results:
                context += f"""
========================================
SOURCE: {filename}
RELEVANCE: {score:.3f}
========================================
{doc[:2000]}...\n"""
        else:
            logger.warning("No relevant templates found - proceeding with generation only")
            context = "[No contract templates found. Generate from first principles.]"
        
        # ======================================
        # STEP 3 — DRAFT CONTRACT
        # ======================================
        
        print("\n[3/5] Drafting initial contract...")
        logger.info("Step 3: Drafting contract")
        
        draft_prompt = f"""You are a senior UK commercial lawyer.

Use the structured specification and any provided source materials to draft a professional legal agreement.

SPECIFICATION:
{structured_spec}

SOURCE MATERIALS:
{context}

INSTRUCTIONS:
1. Draft a complete professional agreement
2. Use formal legal language
3. Maintain consistent defined terms
4. Include logical section headings
5. Optimize for enforceability
6. Make commercially reasonable assumptions when needed"""
        
        contract_text = call_gemini(
            draft_prompt,
            temperature=config.TEMPERATURE,
            max_tokens=config.MAX_OUTPUT_TOKENS
        )
        
        print("\n" + "="*40)
        print(" INITIAL DRAFT")
        print("="*40)
        print(contract_text)
        
        # ======================================
        # STEP 4 — LEGAL REVIEW
        # ======================================
        
        print("\n[4/5] Performing legal review...")
        logger.info("Step 4: Legal review")
        
        review_prompt = f"""You are senior opposing counsel performing a rigorous legal review.

Identify:
- ambiguous language
- inconsistent definitions
- unenforceable provisions
- missing protections
- commercial risks
- drafting weaknesses
- missing clauses
- negotiation vulnerabilities

AGREEMENT:
{contract_text}

Provide detailed review comments."""
        
        review_text = call_gemini(review_prompt, temperature=0.1, max_tokens=4096)
        
        print("\n" + "="*40)
        print(" LEGAL REVIEW")
        print("="*40)
        print(review_text)
        
        # ======================================
        # STEP 5 — REVISE AGREEMENT
        # ======================================
        
        print("\n[5/5] Revising agreement based on review...")
        logger.info("Step 5: Revising agreement")
        
        revision_prompt = f"""You are senior commercial counsel.

Revise this agreement addressing all review comments:

REVIEW COMMENTS:
{review_text}

ORIGINAL AGREEMENT:
{contract_text}

INSTRUCTIONS:
1. Fix all identified issues
2. Improve enforceability and clarity
3. Keep document structure clean
4. Keep defined terms consistent
5. Remove ambiguity
6. Maintain enterprise-grade drafting quality"""
        
        final_contract = call_gemini(
            revision_prompt,
            temperature=0.1,
            max_tokens=config.MAX_OUTPUT_TOKENS
        )
        
        print("\n" + "="*40)
        print(" FINAL AGREEMENT")
        print("="*40)
        print(final_contract)
        
        # ======================================
        # STEP 6 — SAVE OUTPUT
        # ======================================
        
        output_path = save_contract(final_contract)
        
        print("\n" + "="*40)
        print(" COMPLETE")
        print("="*40)
        print(f"\nFinal agreement saved to:\n{output_path}\n")
        logger.info("Contract generation completed successfully")
    
    except Exception as e:
        logger.error(f"Error during contract generation: {e}", exc_info=True)
        print(f"\nERROR: {e}")
        print("Check logs/ directory for details")


if __name__ == "__main__":
    main()
